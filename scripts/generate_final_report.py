import csv
import os
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / ".codex_deps"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


FINAL_DIR = ROOT / "reports" / "final"
SCREENSHOTS_DIR = ROOT / "reports" / "screenshots"
FIGURES_DIR = ROOT / "reports" / "figures"
ARTIFACTS_DIR = ROOT / "artifacts"
DOCX_PATH = FINAL_DIR / "تقرير_مشروع_استرجاع_المعلومات_النهائي.docx"
ARCHITECTURE_PATH = FINAL_DIR / "system_architecture.png"

TEAM = [
    ("الخضر الديواني", "تجهيز Dataset، إدارة qrels، التدريب على Colab، تكامل الملفات النهائية."),
    ("نايا سعدون", "المعالجة المسبقة، معالجة الاستعلام، وتحسين الاستعلام Query Refinement."),
    ("حلا العوض", "TF-IDF، LSA Embedding، وتمثيل الوثائق والاستعلامات."),
    ("ليث ضاهر", "BM25، الاسترجاع الهجين التسلسلي والمتوازي، والترتيب."),
    ("نوال صالح", "التقييم، واجهة Streamlit، RAG، وتجهيز التقرير."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_direction(cell, rtl=True):
    tc_pr = cell._tc.get_or_add_tcPr()
    text_direction = tc_pr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        tc_pr.append(text_direction)
    text_direction.set(qn("w:val"), "rtl" if rtl else "lrTb")


def set_paragraph_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:cs"), "Arial")


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_rtl_paragraph(document, text="", bold=False, size=12, before=0, after=6):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    size = 17 if level == 1 else 14
    color = (31, 78, 121) if level == 1 else (55, 55, 55)
    set_run_font(run, size=size, bold=True, color=color)
    paragraph.style = f"Heading {min(level, 3)}"
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run_font(paragraph.add_run(text), size=11)
    return paragraph


def set_cell_text(cell, text, bold=False, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=10, bold=bold or header, color=(255, 255, 255) if header else None)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_direction(cell)
    if header:
        set_cell_shading(cell, "1F4E79")


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, header=True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def add_caption(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=True, color=(90, 90, 90))
    paragraph.paragraph_format.space_after = Pt(8)


def add_picture(document, path, caption, width=6.3):
    path = Path(path)
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(document, caption)


def read_csv(name):
    with (ARTIFACTS_DIR / name).open(encoding="utf-8", newline="") as file:
        return list(csv.DictReader(file))


def metric_rows(filename, columns):
    rows = []
    for item in read_csv(filename):
        row = [item["method"]]
        row.extend(f"{float(item[column]):.4f}" for column in columns)
        rows.append(row)
    return rows


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("صفحة ")
    set_run_font(run, size=9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_toc(document):
    paragraph = document.add_paragraph()
    set_paragraph_rtl(paragraph)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "يتم تحديث جدول المحتويات تلقائياً عند فتح الملف."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def make_architecture_diagram():
    canvas = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 30)
    bold = ImageFont.truetype(str(bold_path), 32)
    title = ImageFont.truetype(str(bold_path), 42)

    draw.text((520, 30), "Information Retrieval System Architecture", font=title, fill="#17365D")

    boxes = {
        "dataset": (60, 150, 360, 260, "ClinicalTrials Dataset"),
        "data": (480, 150, 780, 260, "Data Service"),
        "db": (920, 90, 1260, 200, "SQLite Raw Document Store"),
        "prep": (920, 250, 1260, 360, "Preprocessing Service"),
        "index": (1400, 250, 1710, 360, "Indexing Service"),
        "models": (1310, 450, 1760, 610, "TF-IDF | BM25 | LSA\nHybrid Serial / Parallel"),
        "retrieval": (820, 500, 1180, 620, "Retrieval & Ranking Service"),
        "refine": (370, 450, 700, 560, "Query Refinement Service"),
        "evaluation": (1280, 740, 1650, 850, "Evaluation Service"),
        "rag": (760, 740, 1110, 850, "Grounded RAG Service"),
        "api": (250, 730, 580, 840, "FastAPI Gateway"),
        "ui": (250, 900, 580, 1010, "Streamlit UI"),
    }

    colors = {
        "dataset": "#E2F0D9",
        "data": "#DDEBF7",
        "db": "#FFF2CC",
        "prep": "#DDEBF7",
        "index": "#DDEBF7",
        "models": "#FCE4D6",
        "retrieval": "#BDD7EE",
        "refine": "#E4DFEC",
        "evaluation": "#E2F0D9",
        "rag": "#F4B183",
        "api": "#C6E0B4",
        "ui": "#C6E0B4",
    }

    for key, (x1, y1, x2, y2, label) in boxes.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=colors[key], outline="#44546A", width=4)
        lines = label.split("\n")
        total_height = len(lines) * 40
        y = (y1 + y2 - total_height) / 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=bold if len(lines) == 1 else font)
            x = (x1 + x2 - (bbox[2] - bbox[0])) / 2
            draw.text((x, y), line, font=bold if len(lines) == 1 else font, fill="#1F1F1F")
            y += 40

    def arrow(start, end):
        draw.line((start, end), fill="#5B9BD5", width=7)
        x2, y2 = end
        x1, y1 = start
        import math

        angle = math.atan2(y2 - y1, x2 - x1)
        size = 18
        left = (x2 - size * math.cos(angle - 0.55), y2 - size * math.sin(angle - 0.55))
        right = (x2 - size * math.cos(angle + 0.55), y2 - size * math.sin(angle + 0.55))
        draw.polygon([end, left, right], fill="#5B9BD5")

    arrow((360, 205), (480, 205))
    arrow((780, 190), (920, 145))
    arrow((780, 225), (920, 305))
    arrow((1260, 305), (1400, 305))
    arrow((1555, 360), (1555, 450))
    arrow((1310, 530), (1180, 560))
    arrow((700, 505), (820, 550))
    arrow((1000, 620), (940, 740))
    arrow((1180, 585), (1280, 785))
    arrow((760, 790), (580, 785))
    arrow((415, 840), (415, 900))
    arrow((920, 145), (580, 785))

    canvas.save(ARCHITECTURE_PATH)


def build_document():
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram()

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    add_page_number(section)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    # Cover
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(80)
    run = paragraph.add_run("مشروع عملي نظم استرجاع المعلومات 2026")
    set_run_font(run, size=24, bold=True, color=(31, 78, 121))

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("بناء نظام استرجاع معلومات طبي")
    set_run_font(run, size=22, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Information Retrieval System")
    set_run_font(run, size=18, bold=True, color=(192, 0, 0))

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(35)
    run = paragraph.add_run("Dataset: ClinicalTrials 2017 / TREC Precision Medicine 2017")
    set_run_font(run, size=14, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("عدد الوثائق: 241,006")
    set_run_font(run, size=14)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(30)
    run = paragraph.add_run("إعداد الفريق")
    set_run_font(run, size=16, bold=True, color=(31, 78, 121))
    for name, _ in TEAM:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(name), size=13)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(35)
    set_run_font(paragraph.add_run("التاريخ: 14 حزيران 2026"), size=12)
    document.add_page_break()

    add_heading(document, "جدول المحتويات", 1)
    add_toc(document)
    document.add_page_break()

    add_heading(document, "1. الملخص التنفيذي", 1)
    add_rtl_paragraph(
        document,
        "يقدم هذا المشروع نظام استرجاع معلومات متخصصاً في التجارب السريرية الطبية. "
        "يقبل النظام استعلام المستخدم باللغة الطبيعية، ويطبق المعالجة المسبقة نفسها المستخدمة للوثائق، "
        "ثم يسترجع ويرتب الوثائق باستخدام TF-IDF وBM25 وLSA Embedding وطرق هجينة تسلسلية ومتوازية. "
        "كما يتضمن Sentence-BERT لإعادة ترتيب مرشحي BM25 وميزة Extractive Grounded RAG لتوليد إجابة "
        "مؤرضة مع الأدلة والمصادر. تم تصميم النظام وفق SOA، ويوفر واجهة Streamlit وREST API باستخدام FastAPI."
    )
    add_rtl_paragraph(
        document,
        "استُخدمت Dataset كاملة من ClinicalTrials 2017 بعدد 241,006 وثيقة و30 استعلاماً و13,019 حكماً "
        "للملاءمة qrels. حُفظ النص الأصلي في SQLite، بينما حُفظت الفهارس والنماذج في ملفات artifacts جاهزة "
        "للتحميل أثناء المقابلة دون إعادة التدريب."
    )

    add_heading(document, "2. أهداف المشروع", 1)
    for item in [
        "بناء محرك بحث قادر على استرجاع أفضل الوثائق الطبية المرتبطة باستعلام المستخدم.",
        "مقارنة تمثيلات TF-IDF وBM25 وEmbedding والتمثيلين الهجينين.",
        "توفير Query Processing وQuery Refinement متوافقين مع تمثيل الوثائق.",
        "تخزين النصوص الأصلية في قاعدة بيانات وقراءتها حسب doc_id وقت الاستعلام.",
        "تقييم النظام باستخدام MAP وnDCG وPrecision وRecall.",
        "تصميم النظام وفق SOA وتوفير واجهة مستخدم وREST API قابلة للاختبار بشكل مستقل.",
        "إضافة RAG كميزة إضافية وتقييم جودة المصادر والـ grounding.",
    ]:
        add_bullet(document, item)

    add_heading(document, "3. توصيف Dataset", 1)
    add_rtl_paragraph(
        document,
        "تم اعتماد Dataset واحدة بناءً على التوضيح النهائي للمعيدة: "
        "clinicaltrials/2017/trec-pm-2017 من مكتبة ir_datasets. وهي لقطة من ClinicalTrials.gov "
        "مستخدمة في مسار TREC 2017 Precision Medicine.",
    )
    add_table(
        document,
        ["الخاصية", "القيمة"],
        [
            ["معرف Dataset", "clinicaltrials/2017/trec-pm-2017"],
            ["عدد الوثائق", "241,006"],
            ["عدد الاستعلامات", "30"],
            ["عدد qrels", "13,019"],
            ["لغة البيانات", "الإنجليزية"],
            ["معرف الوثيقة", "NCT Clinical Trial ID"],
            ["حقول الوثيقة", "title, condition, summary, detailed_description, eligibility"],
            ["حقول الاستعلام", "disease, gene, demographic, other"],
        ],
        [5, 11],
    )
    add_rtl_paragraph(
        document,
        "تم دمج جميع الحقول النصية للوثيقة في حقل نمذجة واحد قبل الفهرسة. أما النص الأصلي الكامل فحُفظ "
        "في قاعدة بيانات SQLite، وعند ظهور النتائج يرجع النظام doc_id ثم يقرأ الوثيقة الأصلية من القاعدة. "
        "يحقق ذلك سرعة في الاسترجاع مع الحفاظ على النص غير المنظف للعرض."
    )

    add_heading(document, "4. مراحل تنفيذ المشروع", 1)
    stages = [
        ("1", "تحميل Dataset", "استخدام ir_datasets وتحميل الوثائق والاستعلامات وqrels."),
        ("2", "تجهيز النص الأصلي", "دمج الحقول وحفظ الوثائق في SQLite على دفعات."),
        ("3", "المعالجة المسبقة", "lowercase، tokenization، إزالة stop words والرموز والكلمات القصيرة."),
        ("4", "الفهرسة", "بناء TF-IDF وBM25 sparse matrix وLSA embeddings على كامل Dataset."),
        ("5", "معالجة الاستعلام", "تطبيق المعالجة نفسها وتمثيل query حسب نموذج البحث."),
        ("6", "الاسترجاع والترتيب", "حساب الدرجات وإرجاع Top K حسب النموذج المختار."),
        ("7", "إعادة الترتيب", "استخدام Sentence-BERT فوق أفضل 50 مرشحاً من BM25."),
        ("8", "RAG", "اختيار أفضل الجمل المرتبطة بالسؤال وإرفاق evidence وcitations."),
        ("9", "التقييم", "حساب المقاييس على qrels وحفظ CSV والرسوم البيانية."),
        ("10", "العرض", "تشغيل Streamlit أو REST API باستخدام artifacts الجاهزة."),
    ]
    add_table(document, ["المرحلة", "الاسم", "الوصف"], stages, [2, 4, 11])

    add_heading(document, "5. المعالجة المسبقة ومعالجة الاستعلام", 1)
    add_rtl_paragraph(
        document,
        "تنفذ TextProcessor عملية موحدة للوثائق والاستعلامات: تحويل الأحرف إلى lowercase، استخراج الكلمات "
        "والأرقام بتعبير منتظم، إزالة stop words الإنجليزية، وإلغاء الرموز والكلمات الأقصر من حرفين. "
        "يمرر النص المنظف إلى TfidfVectorizer باستخدام tokenizer=str.split مع تعطيل lowercase وtoken pattern "
        "الافتراضيين، مما يمنع حدوث تنظيف مزدوج."
    )
    add_heading(document, "5.1 Query Refinement", 2)
    add_rtl_paragraph(
        document,
        "تقوم Query Refinement Service بتصحيح مجموعة من الأخطاء الإملائية الشائعة وإضافة مرادفات محدودة. "
        "جُعلت الميزة اختيارية من الواجهة لأن نتائج التقييم أثبتت أنها تحسن TF-IDF في بعض المقاييس، "
        "لكنها قد تقلل دقة BM25 عند إضافة كلمات عامة إلى استعلامات طبية دقيقة."
    )

    add_heading(document, "6. نماذج تمثيل واسترجاع الوثائق", 1)
    methods = [
        ("TF-IDF", "Vector Space Model", "Cosine similarity بين query vector ووثائق TF-IDF."),
        ("BM25", "Probabilistic model", "Sparse term matrix مع k1 وb قابلين للتغيير من الواجهة."),
        ("LSA Embedding", "TF-IDF + TruncatedSVD", "تمثيل دلالي خفيف بـ64 بعداً للمقارنة الأساسية."),
        ("Hybrid Parallel", "Score fusion", "دمج درجات TF-IDF وBM25 وLSA بعد normalization."),
        ("Hybrid Serial", "Candidate reranking", "BM25 يجلب المرشحين ثم LSA يعيد ترتيبهم."),
        ("Sentence-BERT", "Semantic reranking", "BM25 يجلب 50 مرشحاً ثم all-MiniLM-L6-v2 يعيد ترتيبهم."),
    ]
    add_table(document, ["الطريقة", "النوع", "آلية العمل"], methods, [4, 4, 9])
    add_rtl_paragraph(
        document,
        "يستخدم Hybrid Parallel أوزان دمج افتراضية 0.25 لـTF-IDF و0.60 لـBM25 والباقي لـLSA. "
        "أما Hybrid Serial فيضيف درجة LSA إلى درجة BM25 للمرشحين فقط. يوفر النظام الاختيار بين الطريقتين "
        "من قائمة Retrieval method."
    )

    add_heading(document, "7. الفهرسة والتخزين", 1)
    add_rtl_paragraph(
        document,
        "بُني TF-IDF بحد أقصى 30,000 feature وmin_df=2 وmax_df=0.95 وبنوع float32 لتقليل الذاكرة. "
        "استخدم BM25 مصفوفة sparse مشتركة مع vocabulary الخاص بـTF-IDF. وتم حفظ الفهرس الكامل في "
        "search_index.joblib، بينما بلغ حجم documents.sqlite نحو 1.1GB ويحتوي على 241,006 وثيقة أصلية."
    )
    for item in [
        "documents.sqlite: قاعدة النصوص الأصلية، مفهرسة حسب doc_id.",
        "search_index.joblib: TF-IDF matrix وvectorizer وSVD وBM25 sparse index.",
        "dataset_metadata.json: معلومات Dataset ومعاملات البناء.",
        "evaluation_metrics*.csv: نتائج التقييم الأساسية والإضافية.",
    ]:
        add_bullet(document, item)

    add_heading(document, "8. بنية النظام وفق SOA", 1)
    add_rtl_paragraph(
        document,
        "تم فصل المسؤوليات إلى خدمات مستقلة. ServiceContainer هو نقطة التجميع الوحيدة التي تحمل الفهرس "
        "وقاعدة البيانات وتحقنها في RetrievalService وRagService. تستخدم Streamlit وFastAPI الحاوية نفسها، "
        "وبالتالي لا تعتمد الواجهة على تفاصيل الفهرسة أو التخزين. هذا يحقق Maintainability وReusability "
        "ويتيح اختبار الخدمات دون تشغيل الواجهة."
    )
    add_picture(document, ARCHITECTURE_PATH, "الشكل 1: المخطط المعماري للخدمات وآلية التواصل بينها", 6.5)
    services = [
        ("Data Service", "data_service.py", "تحميل Dataset وqrels ودمج الحقول وتخزين الوثائق."),
        ("Document Store", "database.py", "عمليات SQLite وحفظ النص الأصلي واسترجاعه حسب doc_id."),
        ("Preprocessing", "text_processing.py", "توحيد تنظيف الوثيقة والاستعلام."),
        ("Query Refinement", "query_refinement.py", "التصحيح الإملائي والتوسع بالمرادفات."),
        ("Indexing", "indexing_service.py", "بناء وحفظ TF-IDF وLSA وBM25."),
        ("Retrieval", "retrieval_service.py", "البحث والترتيب والهجين وBERT reranking."),
        ("Evaluation", "evaluation_service.py", "حساب مقاييس IR وتوليد الرسوم."),
        ("RAG", "rag_service.py", "بناء إجابة مؤرضة وإرجاع الأدلة والمصادر."),
        ("API Gateway", "api.py", "REST endpoints: health, search, rag, metrics."),
        ("UI", "app.py", "واجهة Streamlit للتجربة والعرض."),
    ]
    add_table(document, ["الخدمة", "الملف", "المسؤولية"], services, [4, 4, 9])

    add_heading(document, "9. واجهات التشغيل", 1)
    add_heading(document, "9.1 واجهة Streamlit", 2)
    add_rtl_paragraph(
        document,
        "توفر الواجهة اختيار نموذج الاسترجاع، Top K، معاملات BM25 k1 وb، وتشغيل Query Refinement. "
        "كما تحتوي على تبويبات Search وRAG Chat وEvaluation، وتعرض الوثائق الأصلية من SQLite."
    )
    add_heading(document, "9.2 REST API", 2)
    add_table(
        document,
        ["Endpoint", "الوظيفة"],
        [
            ["GET /health", "التحقق من جاهزية النظام وعدد الوثائق."],
            ["POST /search", "البحث مع اختيار النموذج والمعاملات وTop K."],
            ["POST /rag", "إرجاع الإجابة المؤرضة والأدلة والمصادر."],
            ["GET /metrics", "قراءة جميع تقارير التقييم بصيغة JSON."],
        ],
        [5, 12],
    )
    add_rtl_paragraph(
        document,
        "توفر FastAPI توثيق OpenAPI تفاعلياً على http://127.0.0.1:8000/docs، ويمكن اختبار كل endpoint "
        "بشكل مستقل عن Streamlit."
    )

    add_heading(document, "10. الميزة الإضافية: Grounded RAG", 1)
    add_rtl_paragraph(
        document,
        "تسترجع RAG Service أفضل الوثائق بالطريقة المختارة، ثم تحلل جمل كل وثيقة وتختار الجملة التي تحقق "
        "أكبر تغطية لكلمات السؤال. تعرض أفضل ثلاثة أدلة مع أرقام citations، ثم تعرض قائمة المصادر كاملة "
        "مع doc_id ودرجة الاسترجاع. هذا التصميم Extractive RAG، لذلك لا يختلق معلومات خارج الوثائق."
    )

    add_heading(document, "11. منهجية التقييم", 1)
    add_rtl_paragraph(
        document,
        "استخدمت الاستعلامات وqrels الرسمية. استُرجعت أفضل 1000 وثيقة لكل query لحساب MAP@1000 وRecall@1000، "
        "بينما حُسب nDCG@10 وPrecision@10 على الرتب العشر الأولى. تم تقييم 29 استعلاماً لديها qrels."
    )
    for item in [
        "MAP@1000: متوسط الدقة عبر مواضع الوثائق الملائمة حتى عمق 1000.",
        "nDCG@10: جودة ترتيب النتائج العشر الأولى مع مراعاة درجات الملاءمة.",
        "Precision@10: نسبة الوثائق الملائمة ضمن أول عشر نتائج.",
        "Recall@1000: نسبة الوثائق الملائمة التي استرجعها النظام حتى عمق 1000.",
    ]:
        add_bullet(document, item)

    add_heading(document, "12. نتائج التقييم الأساسية", 1)
    base_columns = ["map_at_1000", "ndcg_at_10", "precision_at_10", "recall_at_1000"]
    add_table(
        document,
        ["Method", "MAP@1000", "nDCG@10", "P@10", "Recall@1000"],
        metric_rows("evaluation_metrics.csv", base_columns),
        [4, 3, 3, 3, 3],
    )
    add_rtl_paragraph(
        document,
        "حقق BM25 أفضل MAP@1000 وPrecision@10 وRecall@1000، مما يثبت ملاءمته للمصطلحات الطبية الدقيقة. "
        "حقق Hybrid Serial أفضل nDCG@10، أي أنه رتب الوثائق شديدة الملاءمة جيداً في المراكز الأولى. "
        "كانت LSA منخفضة لأنها ضغطت vocabulary طبية كبيرة إلى 64 بعداً فقط، لذلك استُخدمت baseline ولم "
        "تُعتمد كنموذج دلالي نهائي."
    )
    add_picture(document, FIGURES_DIR / "evaluation_metrics.png", "الشكل 2: مقارنة نماذج الاسترجاع الأساسية", 6.2)

    add_heading(document, "13. نتائج Query Refinement", 1)
    refined_columns = ["map_at_1000", "ndcg_at_10", "precision_at_10", "recall_at_1000"]
    add_table(
        document,
        ["Method", "MAP@1000", "nDCG@10", "P@10", "Recall@1000"],
        metric_rows("evaluation_metrics_refined.csv", refined_columns),
        [4, 3, 3, 3, 3],
    )
    add_rtl_paragraph(
        document,
        "حسنت التوسعة TF-IDF في nDCG@10 وP@10، كما رفعت Recall الخاص بـLSA. لكنها خفضت MAP وRecall في BM25 "
        "والنماذج الهجينة بسبب إدخال مرادفات عامة في استعلامات جينية دقيقة. لذلك بقيت اختيارية."
    )

    add_heading(document, "14. تقييم Sentence-BERT", 1)
    bert_columns = ["map_at_10", "ndcg_at_10", "precision_at_10", "recall_at_10"]
    add_table(
        document,
        ["Method", "MAP@10", "nDCG@10", "P@10", "Recall@10"],
        metric_rows("evaluation_metrics_bert.csv", bert_columns),
        [4, 3, 3, 3, 3],
    )
    add_rtl_paragraph(
        document,
        "حقق BERT nDCG@10=0.2474 وP@10=0.2655، وهي نتائج أعلى بكثير من LSA في أول عشر نتائج. "
        "وقد استُخدم كـreranker فوق BM25 للحفاظ على السرعة وعدم حساب embedding لكل الوثائق وقت query."
    )

    add_heading(document, "15. تقييم RAG قبل وبعد الميزة الإضافية", 1)
    rag = read_csv("rag_evaluation_metrics.csv")[0]
    add_table(
        document,
        ["Source P@5", "Source Recall@5", "Query Coverage", "Citation Coverage", "Groundedness"],
        [[
            f"{float(rag['source_precision_at_5']):.4f}",
            f"{float(rag['source_recall_at_5']):.4f}",
            f"{float(rag['answer_query_coverage']):.4f}",
            f"{float(rag['citation_coverage']):.4f}",
            f"{float(rag['groundedness']):.4f}",
        ]],
        [3, 3, 3, 3, 3],
    )
    add_rtl_paragraph(
        document,
        "قبل RAG يعرض النظام قائمة الوثائق فقط. بعد RAG أصبحت هناك إجابة مختصرة مع evidence وcitations. "
        "Groundedness=1.0 يعني أن جميع جمل الأدلة موجودة في النصوص الأصلية. Citation Coverage=0.6 لأن "
        "الإجابة تستخدم أفضل ثلاثة أدلة من خمسة مصادر مسترجعة."
    )

    add_heading(document, "16. لقطات من النسخة التنفيذية", 1)
    add_picture(document, SCREENSHOTS_DIR / "final_search.png", "الشكل 3: البحث الهجين وعرض الوثيقة الأصلية", 6.4)
    add_picture(document, SCREENSHOTS_DIR / "final_rag.png", "الشكل 4: إجابة Grounded RAG مع citations", 6.4)
    add_picture(
        document,
        SCREENSHOTS_DIR / "final_base_refined_evaluation.png",
        "الشكل 5: التقييم الأساسي وتقييم Query Refinement",
        6.4,
    )
    add_picture(
        document,
        SCREENSHOTS_DIR / "final_bert_rag_evaluation.png",
        "الشكل 6: تقييم Sentence-BERT وRAG",
        6.4,
    )
    add_picture(
        document,
        SCREENSHOTS_DIR / "final_evaluation_chart.png",
        "الشكل 7: الرسم البياني النهائي للمقارنة",
        6.4,
    )

    add_heading(document, "17. الاختبارات وضمان الجودة", 1)
    add_rtl_paragraph(
        document,
        "أضيفت اختبارات مستقلة باستخدام unittest لخدمة SQLite والمعالجة المسبقة وQuery Refinement ومقاييس "
        "التقييم وRAG grounding. كما اختُبرت REST API فعلياً على الفهرس الكامل، واختُبر BM25 وBERT وRAG "
        "من سطر الأوامر. نجحت الاختبارات الأربعة ونجح compile لجميع الملفات."
    )
    add_rtl_paragraph(document, "أمر الاختبار:", bold=True)
    add_rtl_paragraph(document, r".\run_tests.cmd", size=11)

    add_heading(document, "18. تقسيم العمل بين أعضاء الفريق", 1)
    add_table(document, ["العضو", "المهام"], TEAM, [5, 12])

    add_heading(document, "19. النسخة التنفيذية وطريقة التشغيل", 1)
    add_rtl_paragraph(
        document,
        "النسخة التنفيذية جاهزة للاستعلام أثناء المقابلة ولا تحتاج إلى Colab أو إعادة تدريب. يجب أن تبقى "
        "ملفات artifacts الكبيرة على جهاز العرض."
    )
    add_table(
        document,
        ["المكون", "الأمر / الرابط"],
        [
            ["واجهة البحث", r".\run_app.cmd ثم http://localhost:8501"],
            ["REST API", r".\run_api.cmd ثم http://127.0.0.1:8000/docs"],
            ["الاختبارات", r".\run_tests.cmd"],
            ["البحث من CLI", 'python scripts\\search.py "lung cancer EGFR adult" --method bm25 --top-k 10'],
        ],
        [5, 12],
    )
    add_rtl_paragraph(
        document,
        "يوفر المشروع محرك البحث عبر واجهتين تنفيذيتين: واجهة Streamlit للمستخدم، وREST API للتكامل "
        "والاختبار المستقل. ويمكن اختيار TF-IDF أوBM25 أوEmbedding أوHybrid أوBERT من الواجهة."
    )

    add_heading(document, "20. GitHub وبنية المشروع", 1)
    add_rtl_paragraph(document, "رابط GitHub:")
    add_rtl_paragraph(document, "https://github.com/khaderaldiwani/ir_project", bold=True)
    add_table(
        document,
        ["المسار", "المحتوى"],
        [
            ["src/ir_project/services", "الخدمات المستقلة للنظام."],
            ["scripts", "التحضير والتقييم والبحث وتقييم BERT وRAG."],
            ["notebooks", "نوت بوك Colab لإعادة التدريب من الصفر."],
            ["artifacts", "metadata ونتائج التقييم، مع تجاهل الملفات الضخمة في Git."],
            ["reports", "التقرير والصور والرسوم البيانية."],
            ["app.py", "واجهة Streamlit."],
            ["api.py", "FastAPI Gateway."],
            ["README.md", "شرح البنية والتشغيل والتقييم."],
        ],
        [6, 11],
    )

    add_heading(document, "21. القيود والتحسينات المستقبلية", 1)
    for item in [
        "استخدام Domain-specific biomedical Sentence-BERT قد يحسن النتائج الدلالية أكثر.",
        "يمكن استبدال Extractive RAG بمولد لغوي محلي مع الحفاظ على citations.",
        "يمكن نقل SQLite إلى PostgreSQL أو MongoDB عند التشغيل الموزع.",
        "يمكن إضافة Learning to Rank لتعلم أوزان الدمج بدلاً من تثبيتها يدوياً.",
        "يمكن إضافة cache لنتائج BERT المتكررة ومراقبة زمن الاستجابة.",
    ]:
        add_bullet(document, item)

    add_heading(document, "22. الخلاصة", 1)
    add_rtl_paragraph(
        document,
        "تم بناء نظام استرجاع معلومات متكامل على Dataset كاملة تتجاوز 200 ألف وثيقة. يحقق النظام جميع "
        "المراحل المطلوبة من المعالجة والفهرسة والتمثيل والاسترجاع الهجين وتحسين الاستعلام والتقييم. "
        "كما يطبق SOA عملياً عبر خدمات مستقلة وServiceContainer وFastAPI، ويوفر نسخة تنفيذية مستقرة "
        "للواجهة وAPI. أظهرت النتائج تفوق BM25 وHybrid Serial، وأثبت BERT فائدته كـreranker، بينما حقق "
        "RAG groundedness كاملاً مع عرض مصادر واضحة."
    )

    add_heading(document, "23. المصادر والمراجع", 1)
    references = [
        "IR Datasets, Clinical Trials collection: https://ir-datasets.com/clinicaltrials.html",
        "Roberts, K. et al. Overview of the TREC 2017 Precision Medicine Track, TREC 2017.",
        "TREC Precision Medicine data: https://trec.nist.gov/data/precmed.html",
        "Robertson, S. and Zaragoza, H. The Probabilistic Relevance Framework: BM25 and Beyond, 2009. https://doi.org/10.1561/1500000019",
        "scikit-learn TfidfVectorizer: https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html",
        "scikit-learn TruncatedSVD: https://scikit-learn.org/stable/modules/generated/sklearn.decomposition.TruncatedSVD.html",
        "Reimers, N. and Gurevych, I. Sentence-BERT, 2019: https://arxiv.org/abs/1908.10084",
        "Sentence Transformers documentation: https://www.sbert.net/",
        "FastAPI documentation: https://fastapi.tiangolo.com/",
        "Streamlit documentation: https://docs.streamlit.io/",
        "SQLite documentation: https://www.sqlite.org/docs.html",
    ]
    for index, reference in enumerate(references, start=1):
        add_rtl_paragraph(document, f"{index}. {reference}", size=10)

    document.core_properties.title = "تقرير مشروع نظم استرجاع المعلومات 2026"
    document.core_properties.subject = "ClinicalTrials Information Retrieval System"
    document.core_properties.author = "فريق مشروع نظم استرجاع المعلومات"
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
